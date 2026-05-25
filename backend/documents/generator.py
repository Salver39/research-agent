"""Generate .docx research documents from session state."""

from __future__ import annotations

import os
from pathlib import Path
from typing import Any, Optional

from docx import Document as DocxDocument

from agents.base import selected_hypotheses


def _doc_definitions() -> dict:
    return {
        "Discussion Guide":         _guide,
        "Рекрутинговый скринер":   _screener,
        "Briefing для интервьюера": _briefing,
        "Шаблон инсайтов":         _insights,
        "Дизайн исследования":     _design_doc,
        "Чеклист запуска":         _checklist,
    }


async def generate_all(
    state: dict[str, Any],
    output_dir: str,
    doc_name: Optional[str] = None,
    fmt: str = "docx",
) -> Any:
    defs = _doc_definitions()
    if doc_name:
        fn = defs.get(doc_name)
        if fn is None:
            return None
        return _save(fn(state), output_dir, doc_name, fmt)

    paths: list[str] = []
    for name, fn in defs.items():
        doc = fn(state)
        docx_path = _save(doc, output_dir, name, "docx")
        if docx_path:
            paths.append(docx_path)
        pdf_path = _save(doc, output_dir, name, "pdf")
        if pdf_path:
            paths.append(pdf_path)
    return paths


def _save(doc: DocxDocument, output_dir: str, name: str, fmt: str) -> Optional[str]:
    import re as _re
    safe = _re.sub(r"[^\w\s\-.]", "", name).replace(" ", "_")
    docx_path = os.path.join(output_dir, f"{safe}.docx")
    doc.save(docx_path)

    if fmt == "pdf":
        pdf_path = os.path.join(output_dir, f"{safe}.pdf")
        try:
            import mammoth
            from weasyprint import HTML
            with open(docx_path, "rb") as f:
                html_content = mammoth.convert_to_html(f).value
            HTML(string=html_content).write_pdf(pdf_path)
        except Exception:
            return None
        return pdf_path

    return docx_path


# ── Helpers ───────────────────────────────────────────────────────────────────

def _h1(doc: DocxDocument, text: str):
    p = doc.add_heading(text, level=1)
    return p

def _h2(doc: DocxDocument, text: str):
    p = doc.add_heading(text, level=2)
    return p

def _meta(doc: DocxDocument, label: str, value: str):
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(value or "—")


# ── Document builders ─────────────────────────────────────────────────────────

def _guide(state: dict) -> DocxDocument:
    doc = DocxDocument()
    brief = state.get("brief") or {}
    method = state.get("method") or {}
    design = state.get("design") or {}

    _h1(doc, "Discussion Guide")
    _meta(doc, "Тема", brief.get("research_question", "—"))
    _meta(doc, "Метод", method.get("name", "—"))
    _meta(doc, "Длительность", design.get("total_duration", "—"))
    _meta(doc, "Ключевое решение", brief.get("decision", "—"))
    doc.add_paragraph()

    # Usability testing format
    if design.get("tasks"):
        pre = design.get("pre_interview") or {}
        if pre:
            _h2(doc, "Преинтервью")
            if pre.get("goal"):
                doc.add_paragraph(f"Цель: {pre['goal']}")
            for q in (pre.get("questions") or []):
                doc.add_paragraph(f"— {q}", style="List Bullet")
            doc.add_paragraph()

        for i, task in enumerate(design["tasks"], 1):
            title = task.get("title") or f"Задание {i}"
            _h2(doc, f"Задание {i}. {title}")
            if task.get("hypothesis_text"):
                p = doc.add_paragraph()
                run = p.add_run("Гипотеза: ")
                run.bold = True
                p.add_run(task["hypothesis_text"])
            if task.get("scenario"):
                p = doc.add_paragraph()
                run = p.add_run("Легенда: ")
                run.bold = True
                p.add_run(task["scenario"])
            if task.get("task"):
                p = doc.add_paragraph()
                run = p.add_run("Задание: ")
                run.bold = True
                p.add_run(task["task"])
            if task.get("observe"):
                p = doc.add_paragraph()
                run = p.add_run("Наблюдаем: ")
                run.bold = True
                p.add_run(task["observe"])
            if task.get("followup"):
                doc.add_paragraph().add_run("Follow-up вопросы:").bold = True
                for q in task["followup"]:
                    doc.add_paragraph(f"— {q}", style="List Bullet")
            if task.get("success_criteria"):
                p = doc.add_paragraph()
                run = p.add_run("Критерий успеха: ")
                run.bold = True
                p.add_run(task["success_criteria"])
            doc.add_paragraph()

        sus = design.get("sus")
        if sus:
            _h2(doc, "SUS-шкала (опционально)")
            doc.add_paragraph(sus.get("scale", ""))
            for s in (sus.get("statements") or []):
                doc.add_paragraph(f"☐  {s}")
            doc.add_paragraph()

        closing = design.get("closing") or {}
        if closing.get("questions"):
            _h2(doc, "Завершение")
            for q in closing["questions"]:
                doc.add_paragraph(f"— {q}", style="List Bullet")

    # Survey format (main_blocks + screener/warmup/demographics)
    elif design.get("main_blocks"):
        _render_survey(doc, design)

    # Interview format (guide_blocks with optional meta + intro)
    elif design.get("guide_blocks"):
        meta = design.get("meta") or {}
        if meta:
            _h2(doc, "Методологическая шапка")
            if meta.get("goal"):
                p = doc.add_paragraph()
                p.add_run("Цель: ").bold = True
                p.add_run(meta["goal"])
            if meta.get("tasks"):
                doc.add_paragraph().add_run("Задачи:").bold = True
                for t in meta["tasks"]:
                    doc.add_paragraph(f"• {t}", style="List Bullet")
            if meta.get("audience"):
                p = doc.add_paragraph()
                p.add_run("Целевая аудитория: ").bold = True
                p.add_run(meta["audience"])
            doc.add_paragraph()

        intro = design.get("intro") or {}
        if intro.get("items"):
            _h2(doc, f"Вступление ({intro.get('duration', '5 мин')})")
            for item in intro["items"]:
                doc.add_paragraph(f"• {item}", style="List Bullet")
            doc.add_paragraph()

        for block in design["guide_blocks"]:
            _h2(doc, f"{block.get('title', '—')} ({block.get('duration', '?')})")
            if block.get("goal"):
                doc.add_paragraph(f"Цель: {block['goal']}")
            for q in (block.get("questions") or []):
                doc.add_paragraph(f"— {q}", style="List Bullet")
            if block.get("probes"):
                doc.add_paragraph().add_run("Зондирующие вопросы:").bold = True
                for p in block["probes"]:
                    doc.add_paragraph(f"↳ {p}", style="List Bullet")
            doc.add_paragraph()

        closing = design.get("closing") or {}
        if closing.get("questions"):
            _h2(doc, "Завершение")
            for q in closing["questions"]:
                doc.add_paragraph(f"— {q}", style="List Bullet")

    else:
        # Fallback template
        for title, dur, desc in [
            ("Разогрев и контекст", "5 мин", "Снять напряжение, понять бэкграунд участника"),
            ("Текущее поведение", "10 мин", "Как сейчас решает задачу"),
            ("Боли и барьеры", "15 мин", "Где спотыкается, что раздражает"),
            ("Проверка гипотез", "15 мин", "Целевые вопросы под каждую гипотезу"),
            ("Реакция на решение", "10 мин", "Показать концепт / продукт"),
            ("Завершение", "5 мин", "Финальный вопрос, благодарность"),
        ]:
            _h2(doc, f"{title} ({dur})")
            doc.add_paragraph(desc)

    return doc


_SURVEY_TYPE_LABELS = {
    "single_choice": "Одиночный выбор",
    "multi_choice":  "Множественный выбор",
    "likert_5":      "Шкала 1–5",
    "likert_7":      "Шкала 1–7",
    "ranking":       "Ранжирование",
    "numeric":       "Число",
    "open":          "Открытый",
}


def _render_survey_question(doc: DocxDocument, q: dict, show_qualifying: bool = False):
    qid = q.get("id", "")
    qtext = q.get("text", "")
    qtype = q.get("type", "")
    type_label = _SURVEY_TYPE_LABELS.get(qtype, qtype)

    p = doc.add_paragraph()
    run = p.add_run(f"{qid}. ")
    run.bold = True
    p.add_run(qtext)
    if q.get("required"):
        p.add_run(" *")

    meta_bits = [f"Тип: {type_label}"]
    if q.get("measurement_type"):
        meta_bits.append(f"измерение: {q['measurement_type']}")
    if q.get("time_period"):
        meta_bits.append(f"период: {q['time_period']}")
    if q.get("unit"):
        meta_bits.append(f"единица: {q['unit']}")
    if q.get("hypothesis_ids"):
        meta_bits.append(f"гипотезы: {', '.join(q['hypothesis_ids'])}")
    p = doc.add_paragraph()
    p.add_run("    " + " · ".join(meta_bits)).italic = True

    for opt in (q.get("options") or []):
        doc.add_paragraph(f"    ○ {opt}")

    scale = q.get("scale") or {}
    if scale:
        doc.add_paragraph(f"    Шкала: {scale.get('type', '')}")
        for k, v in (scale.get("labels") or {}).items():
            doc.add_paragraph(f"        {k} — {v}")

    if show_qualifying and q.get("qualifying_answers"):
        p = doc.add_paragraph()
        run = p.add_run("    ✓ Проходит дальше: ")
        run.bold = True
        p.add_run(", ".join(q["qualifying_answers"]))

    if q.get("randomize_options"):
        p = doc.add_paragraph()
        p.add_run("    (варианты рандомизируются)").italic = True

    doc.add_paragraph()


def _render_survey(doc: DocxDocument, design: dict):
    meta = design.get("meta") or {}
    if meta:
        _h2(doc, "Методологическая шапка")
        if meta.get("goal"):
            p = doc.add_paragraph()
            p.add_run("Цель опроса: ").bold = True
            p.add_run(meta["goal"])
        if meta.get("tasks"):
            doc.add_paragraph().add_run("Задачи:").bold = True
            for t in meta["tasks"]:
                doc.add_paragraph(f"• {t}", style="List Bullet")
        if meta.get("audience"):
            p = doc.add_paragraph()
            p.add_run("Целевая аудитория: ").bold = True
            p.add_run(meta["audience"])
        if meta.get("estimated_time"):
            p = doc.add_paragraph()
            p.add_run("Время прохождения: ").bold = True
            p.add_run(meta["estimated_time"])
        doc.add_paragraph()

    intro = design.get("intro") or {}
    if intro.get("items"):
        _h2(doc, intro.get("title") or "Вступительный экран")
        for item in intro["items"]:
            doc.add_paragraph(f"• {item}", style="List Bullet")
        doc.add_paragraph()

    screener = design.get("screener") or {}
    if screener.get("questions"):
        _h2(doc, "Скринер")
        if screener.get("goal"):
            doc.add_paragraph(f"Цель: {screener['goal']}")
        for q in screener["questions"]:
            _render_survey_question(doc, q, show_qualifying=True)

    warmup = design.get("warmup") or {}
    if warmup.get("questions"):
        _h2(doc, "Разогрев")
        if warmup.get("goal"):
            doc.add_paragraph(f"Цель: {warmup['goal']}")
        for q in warmup["questions"]:
            _render_survey_question(doc, q)

    for block in design.get("main_blocks") or []:
        _h2(doc, block.get("title") or "Основной блок")
        if block.get("hypothesis_text"):
            p = doc.add_paragraph()
            run = p.add_run("Гипотеза: ")
            run.bold = True
            p.add_run(block["hypothesis_text"])
        for q in (block.get("questions") or []):
            _render_survey_question(doc, q)

    open_qs = design.get("open_questions") or []
    if open_qs:
        _h2(doc, "Открытые вопросы")
        for q in open_qs:
            _render_survey_question(doc, q)

    demo = design.get("demographics") or []
    if demo:
        _h2(doc, "Демография")
        for q in demo:
            _render_survey_question(doc, q)

    routing = design.get("routing") or []
    if routing:
        _h2(doc, "Маршрутизация (skip logic)")
        for r in routing:
            ans = " / ".join(r.get("if_answer_in") or [])
            line = f"Если {r.get('if_question', '?')} = {ans} → {r.get('skip_to', '?')}"
            if r.get("reason"):
                line += f" ({r['reason']})"
            doc.add_paragraph(f"• {line}", style="List Bullet")
        doc.add_paragraph()

    closing = design.get("closing") or {}
    if closing.get("items"):
        _h2(doc, closing.get("title") or "Финальный экран")
        for item in closing["items"]:
            doc.add_paragraph(f"• {item}", style="List Bullet")


def _screener(state: dict) -> DocxDocument:
    doc = DocxDocument()
    sample = state.get("sample") or {}
    method = state.get("method") or {}

    _h1(doc, "Рекрутинговый скринер")
    _meta(doc, "Метод", method.get("name", "—"))
    _meta(doc, "Всего участников", str(sample.get("total_size", "—")))
    doc.add_paragraph()

    segments = sample.get("segments") or []
    if segments:
        _h2(doc, "Сегменты участников")
        for seg in segments:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f"{seg.get('name', '—')} ({seg.get('size', '?')} чел.): ")
            run.bold = True
            p.add_run(seg.get("description", ""))
        doc.add_paragraph()

    criteria = sample.get("criteria") or {}
    if criteria.get("include"):
        _h2(doc, "Критерии включения")
        for c in criteria["include"]:
            doc.add_paragraph(f"✓ {c}", style="List Bullet")
    if criteria.get("exclude"):
        _h2(doc, "Критерии исключения")
        for c in criteria["exclude"]:
            doc.add_paragraph(f"✗ {c}", style="List Bullet")

    screener_qs = sample.get("screener") or []
    if screener_qs:
        doc.add_paragraph()
        _h2(doc, "Скринер-вопросы")
        for i, q in enumerate(screener_qs, 1):
            doc.add_paragraph(f"{i}. {q}")
            doc.add_paragraph("Ответ: ________________________________")
            doc.add_paragraph()

    return doc


def _briefing(state: dict) -> DocxDocument:
    doc = DocxDocument()
    brief = state.get("brief") or {}
    method = state.get("method") or {}
    hypotheses = selected_hypotheses(state)

    _h1(doc, "Briefing для интервьюера")
    _meta(doc, "Метод", method.get("name", "—"))
    _meta(doc, "Длительность", method.get("duration", "—"))
    _meta(doc, "Формат", method.get("format", "—"))
    doc.add_paragraph()

    _h2(doc, "Цель исследования")
    doc.add_paragraph(brief.get("research_question", "—"))
    doc.add_paragraph()

    _h2(doc, "Что важно проверить")
    for h in hypotheses:
        if not isinstance(h, dict):
            doc.add_paragraph(f"• {h}", style="List Bullet")
            continue
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(h.get("text", "")).bold = False
        details = []
        if h.get("verification_method"):
            details.append(f"Проверка: {h['verification_method']}")
        if h.get("action_if_confirmed"):
            details.append(f"Если подтвердится: {h['action_if_confirmed']}")
        if details:
            doc.add_paragraph("    " + " · ".join(details))
    doc.add_paragraph()

    _h2(doc, "Правила проведения")
    rules = [
        "Держитесь нейтрально — не подсказывайте ответы и не реагируйте оценочно.",
        "Давайте паузы после вопроса — участник должен заполнять тишину.",
        "Задавайте уточняющие вопросы: «Расскажите подробнее», «Почему?», «Что произошло дальше?».",
        "Не произносите слова «правильно», «отлично», «хорошо» — это создаёт bias.",
        "Записывайте дословные цитаты, а не интерпретации.",
    ]
    for r in rules:
        doc.add_paragraph(r, style="List Bullet")

    return doc


def _insights(state: dict) -> DocxDocument:
    doc = DocxDocument()
    hypotheses = selected_hypotheses(state)

    _h1(doc, "Шаблон записи инсайтов")
    doc.add_paragraph("Участник: ___________________________")
    doc.add_paragraph("Дата и время: ___________________________")
    doc.add_paragraph("Интервьюер: ___________________________")
    doc.add_paragraph()

    for h in hypotheses:
        text = h.get("text", "") if isinstance(h, dict) else str(h)
        h_id = h.get("id", "") if isinstance(h, dict) else ""
        _h2(doc, f"[{h_id}] {text[:100]}")
        if isinstance(h, dict) and h.get("verification_method"):
            p = doc.add_paragraph()
            p.add_run("Метод проверки: ").bold = True
            p.add_run(h["verification_method"])
        doc.add_paragraph("Результат: Подтверждена  /  Опровергнута  /  Требует доп. проверки")
        doc.add_paragraph("Наблюдения:")
        doc.add_paragraph("    ____________________________________________________")
        doc.add_paragraph("Цитаты:")
        doc.add_paragraph("    ____________________________________________________")
        doc.add_paragraph()

    _h2(doc, "Другие наблюдения")
    doc.add_paragraph("________________________________________________________")
    doc.add_paragraph("________________________________________________________")

    return doc


def _design_doc(state: dict) -> DocxDocument:
    doc = DocxDocument()
    brief = state.get("brief") or {}
    method = state.get("method") or {}
    sample = state.get("sample") or {}
    design = state.get("design") or {}
    hypotheses = selected_hypotheses(state)

    _h1(doc, "Дизайн исследования")
    doc.add_paragraph()

    _h2(doc, "Исследовательский вопрос")
    doc.add_paragraph(brief.get("research_question", "—"))

    _h2(doc, "Ключевое решение")
    doc.add_paragraph(brief.get("decision", "—"))

    _h2(doc, "Метод")
    doc.add_paragraph(f"{method.get('name', '—')} — {method.get('rationale', '')}")
    doc.add_paragraph(f"Участников: {method.get('participants', '?')} | Длительность: {method.get('duration', '?')} | {method.get('format', '')}")

    _h2(doc, "Выборка")
    doc.add_paragraph(f"Всего: {sample.get('total_size', '?')} участников")
    for seg in (sample.get("segments") or []):
        doc.add_paragraph(f"• {seg.get('name', '—')}: {seg.get('size', '?')} чел.", style="List Bullet")

    _h2(doc, "Гипотезы")
    for h in hypotheses:
        text = h.get("text", "") if isinstance(h, dict) else str(h)
        h_id = h.get("id", "") if isinstance(h, dict) else ""
        doc.add_paragraph(f"[{h_id}] {text}", style="List Bullet")

    _h2(doc, "Ограничения")
    doc.add_paragraph(brief.get("constraints", "—"))

    qc = design.get("quality_checks") or {}
    if qc:
        _h2(doc, "Проверка качества")
        labels = {
            "has_decision": "Есть decision",
            "method_matches_uncertainty": "Метод обоснован",
            "hypotheses_covered": "Гипотезы покрыты",
            "no_leading_questions": "Нет leading questions",
        }
        for key, ok in qc.items():
            mark = "✓" if ok else "✗"
            doc.add_paragraph(f"{mark} {labels.get(key, key)}")

    return doc


def _checklist(state: dict) -> DocxDocument:
    doc = DocxDocument()
    _h1(doc, "Чеклист запуска исследования")
    doc.add_paragraph()

    sections = {
        "Подготовка": [
            "Бриф подтверждён заказчиком",
            "Гипотезы проверены на фальсифицируемость",
            "Метод соответствует типу неопределённости",
            "Discussion guide готов и проверен",
        ],
        "Рекрутинг": [
            "Скринер разослан и заполнен",
            "Выборка подобрана по критериям",
            "Участники подтвердили участие",
            "Ссылки на интервью отправлены",
        ],
        "Проведение": [
            "Briefing для интервьюера прочитан",
            "Шаблон инсайтов подготовлен",
            "Запись включена (с согласия участника)",
            "Тихое место и стабильный интернет",
        ],
    }
    for section, items in sections.items():
        _h2(doc, section)
        for item in items:
            doc.add_paragraph(f"☐  {item}")
        doc.add_paragraph()

    return doc
